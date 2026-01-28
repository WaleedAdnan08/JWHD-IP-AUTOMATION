import io
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any

from app.models.office_action import OfficeActionExtractedData, Rejection, ClaimStatus, Objection, ExaminerStatement

logger = logging.getLogger(__name__)

class ReportGenerator:
    def generate_office_action_report(self, data: Dict[str, Any]) -> io.BytesIO:
        """
        Generates a Word document report for a Patent Office Action.
        """
        # Convert dict to model if necessary
        if isinstance(data, dict):
            try:
                oa_data = OfficeActionExtractedData(**data)
            except Exception as e:
                logger.error(f"Failed to parse data into model: {e}")
                raise ValueError("Invalid Office Action Data format")
        else:
            oa_data = data

        document = Document()
        
        # --- TITLE ---
        title = document.add_heading('Office Action Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- SECTION 1: APPLICATION SUMMARY ---
        document.add_heading('1. Application Summary', level=1)
        
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Details'
        
        # Data Rows
        header_info = [
            ("Application Number", oa_data.header.application_number),
            ("Filing Date", oa_data.header.filing_date or "N/A"),
            ("Office Action Date", oa_data.header.office_action_date),
            ("Type", oa_data.header.office_action_type),
            ("Examiner", oa_data.header.examiner_name or "N/A"),
            ("Examiner Type", oa_data.header.examiner_type or "N/A"),
            ("Examiner Phone", oa_data.header.examiner_phone or "N/A"),
            ("Examiner Email", oa_data.header.examiner_email or "N/A"),
            ("Art Unit", oa_data.header.art_unit or "N/A"),
            ("Response Deadline", oa_data.header.response_deadline or "Unknown"),
            ("Confirmation No.", oa_data.header.confirmation_number or "N/A"),
            ("Attorney Docket No.", oa_data.header.attorney_docket_number or "N/A"),
            ("Customer Number", oa_data.header.customer_number or "N/A"),
            ("Title of Invention", oa_data.header.title_of_invention or "N/A"),
            ("First Named Inventor", oa_data.header.first_named_inventor or "N/A"),
            ("Applicant Name", oa_data.header.applicant_name or "N/A"),
        ]
        
        for key, value in header_info:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
            
        document.add_paragraph() # Spacer

        # --- SECTION 2: CLAIMS STATUS OVERVIEW ---
        document.add_heading('2. Claims Status Overview', level=1)
        
        if oa_data.claims_status:
            claims_table = document.add_table(rows=1, cols=3)
            claims_table.style = 'Table Grid'
            hdr_cells = claims_table.rows[0].cells
            hdr_cells[0].text = 'Claim No.'
            hdr_cells[1].text = 'Status'
            hdr_cells[2].text = 'Type'
            
            for claim in oa_data.claims_status:
                row_cells = claims_table.add_row().cells
                row_cells[0].text = claim.claim_number
                row_cells[1].text = claim.status
                row_cells[2].text = claim.dependency_type
        else:
            document.add_paragraph("No specific claim status information extracted.")

        document.add_paragraph()

        # --- SECTION 3: DETAILED REJECTION ANALYSIS ---
        document.add_heading('3. Detailed Rejection Analysis', level=1)
        
        if not oa_data.rejections:
             document.add_paragraph("No rejections found in this Office Action.")
        
        for i, rejection in enumerate(oa_data.rejections):
            # Rejection Header
            heading = document.add_heading(f'Rejection #{i+1}: {rejection.rejection_type}', level=2)
            
            # Details
            p = document.add_paragraph()
            p.add_run("Claims Affected: ").bold = True
            p.add_run(", ".join(rejection.affected_claims))
            
            p = document.add_paragraph()
            p.add_run("Statutory Basis: ").bold = True
            p.add_run(rejection.statutory_basis or "Not specified")

            # Prior Art
            document.add_heading('Prior Art Cited:', level=3)
            if rejection.cited_prior_art:
                for art in rejection.cited_prior_art:
                    p = document.add_paragraph(style='List Bullet')
                    txt = f"{art.identifier}"
                    if art.title:
                        txt += f" - {art.title}"
                    p.add_run(txt)
            else:
                document.add_paragraph("No specific prior art references cited for this rejection.")

            # Reasoning
            document.add_heading('Examiner\'s Reasoning:', level=3)
            document.add_paragraph(rejection.examiner_reasoning)
            
            # Separator
            document.add_paragraph('_' * 40)

        # --- SECTION 4: OBJECTIONS ---
        if oa_data.objections:
            document.add_heading('4. Objections', level=1)
            for i, obj in enumerate(oa_data.objections):
                document.add_heading(f"Objection #{i+1}: {obj.objected_item}", level=2)
                p = document.add_paragraph()
                p.add_run("Reason: ").bold = True
                p.add_run(obj.reason)
                if obj.corrective_action:
                    p = document.add_paragraph()
                    p.add_run("Required Action: ").bold = True
                    p.add_run(obj.corrective_action)

        # --- SECTION 5: EXAMINER STATEMENTS ---
        if oa_data.other_statements:
            document.add_heading('5. Examiner Comments & Allowable Subject Matter', level=1)
            for stmt in oa_data.other_statements:
                document.add_heading(stmt.statement_type, level=2)
                document.add_paragraph(stmt.content)
        
        # --- SECTION 6: DEADLINE ---
        document.add_heading('6. Response Deadline', level=1)
        p = document.add_paragraph()
        p.add_run("Calculated Deadline: ").bold = True
        p.add_run(oa_data.header.response_deadline or "Please verify manually based on mailing date.")
        
        # Save to IO stream
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        
        return output

report_generator = ReportGenerator()